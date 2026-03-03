from typing import Any

from docx import Document as DocxDocument

from app.services.chunking.base import ChunkResult, ChunkingStrategy
from app.utils.text import clean_text, count_tokens

TARGET_TOKENS = 600
MAX_TOKENS = 800
MIN_TOKENS = 50

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


class DocxChunker(ChunkingStrategy):
    """Chunks Word documents by heading hierarchy."""

    async def chunk(self, content: Any, metadata: dict[str, Any]) -> list[ChunkResult]:
        """Content should be a file path to the .docx file."""
        filepath: str = content
        doc_title = metadata.get("title", "")
        doc = DocxDocument(filepath)

        sections = _extract_sections(doc)
        chunks: list[ChunkResult] = []

        for heading, body in sections:
            sub_parts = _split_by_token_limit(body, TARGET_TOKENS, MAX_TOKENS)

            for j, text in enumerate(sub_parts):
                if count_tokens(text) < MIN_TOKENS:
                    continue

                header = self._build_context_header(
                    doc_title,
                    section=heading,
                    prev_topic=chunks[-1].topic_label if chunks else None,
                )

                chunks.append(ChunkResult(
                    content=clean_text(f"[{header}]\n\n{text}"),
                    chunk_index=len(chunks),
                    chunk_type=_detect_type(text),
                    section_heading=heading,
                    metadata={"section": heading, "sub_index": j},
                ))

        return chunks


def _extract_sections(doc: DocxDocument) -> list[tuple[str, str]]:
    """Extract (heading, body) pairs from a Word document."""
    sections: list[tuple[str, str]] = []
    current_heading = ""
    current_body: list[str] = []

    for para in doc.paragraphs:
        if para.style.name in HEADING_STYLES:
            if current_body:
                sections.append((current_heading, "\n\n".join(current_body)))
                current_body = []
            current_heading = para.text.strip()
        else:
            text = para.text.strip()
            if text:
                current_body.append(text)

    # Process tables as structured text
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            current_body.append("\n".join(rows))

    if current_body:
        sections.append((current_heading, "\n\n".join(current_body)))

    return sections


def _split_by_token_limit(text: str, target: int, maximum: int) -> list[str]:
    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current: list[str] = []
    current_tokens = 0

    for para in paragraphs:
        para_tokens = count_tokens(para)
        if current_tokens + para_tokens > maximum and current:
            chunks.append("\n\n".join(current))
            current = [para]
            current_tokens = para_tokens
        else:
            current.append(para)
            current_tokens += para_tokens

    if current:
        chunks.append("\n\n".join(current))
    return chunks


def _detect_type(text: str) -> str:
    lower = text.lower()
    if "definition" in lower or "definiert" in lower:
        return "definition"
    if "beispiel" in lower or "example" in lower:
        return "example"
    if "|" in text and text.count("|") > 3:
        return "overview"
    return "theory"
